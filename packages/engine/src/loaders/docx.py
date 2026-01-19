"""
DOCX file loader using python-docx.
"""

from pathlib import Path
from .base import BaseLoader, LoadResult


class DocxLoader(BaseLoader):
    """Loader for DOCX files"""

    extensions = ['.docx']
    language = "text"

    def load(self, file_path: Path) -> LoadResult:
        """Load text content from DOCX"""
        try:
            from docx import Document
        except ImportError:
            return LoadResult(
                content="",
                language="text",
                metadata={**self.get_metadata(file_path), "error": "python-docx not installed"}
            )

        try:
            doc = Document(file_path)
            paragraphs = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            content = "\n\n".join(paragraphs)

            metadata = self.get_metadata(file_path)
            metadata['paragraphs'] = len(paragraphs)

            return LoadResult(
                content=content,
                language="text",
                metadata=metadata
            )

        except Exception as e:
            return LoadResult(
                content="",
                language="text",
                metadata={**self.get_metadata(file_path), "error": str(e)}
            )
